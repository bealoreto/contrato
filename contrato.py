from docx import Document
from datetime import datetime

doc = Document()
doc.add_heading('CONTRATO DE PRESTAÇÃO DE SERVIÇOS DE CRIAÇÃO DE MÍDIAS DIGITAIS', level=1)

contrato_texto = [
    ("CONTRATANTE:", "Nome: [Nome do cliente]\nCPF/CNPJ: [Número]\nEndereço: [Endereço completo]\nE-mail: [E-mail]"),
    ("CONTRATADO:", "Nome: [Seu nome ou nome da empresa]\nCPF/CNPJ: [Número]\nEndereço: [Endereço completo]\nE-mail: [E-mail]"),
    ("CLÁUSULA 1 – DO OBJETO", 
     "1.1 O presente contrato tem como objeto a criação de mídias digitais, incluindo, mas não se limitando a:\n"
     "[Ex: Todo o processo de Criação de Conteúdo digital, 2 diárias de produção midiática gerando: 8 vídeos para Reels/mês, 1story/dia, etc.]\n\n"
     "1.2 O material será planejado e desenvolvido conforme necessidade do cliente."),
    ("CLÁUSULA 2 – DO PRAZO E ENTREGA", 
     "2.1 O serviço terá início em [data] e término previsto para [data], podendo ser prorrogado mediante acordo por escrito.\n\n"
     "2.2 As entregas serão realizadas conforme cronograma acordado entre as partes."),
    ("CLÁUSULA 3 – DO VALOR E FORMA DE PAGAMENTO", 
     "3.1 O CONTRATANTE pagará ao CONTRATADO o valor de R$1000,00 + 10% do faturamento bruto da empresa, conforme as condições abaixo:\n"
     "- Forma de pagamento: [Pix, transferência bancária, etc.]\n"
     "- Vencimentos: Todo dia 30, o CONTRATANTE irá reportar as vendas do mês ao CONTRATADO, para que seja feito o cálculo do valor do pagamento, que será realizado no dia 5 do mes subsequente à venda\n\n"
    ("CLÁUSULA 4 – DOS DIREITOS AUTORAIS E USO", 
     "4.1 Os direitos de uso das mídias criadas serão transferidos ao CONTRATANTE após o pagamento integral do serviço.\n\n"
     "4.2 O CONTRATADO poderá utilizar as peças criadas em seu portfólio, salvo se houver cláusula de confidencialidade (ver cláusula 6)."),
    ("CLÁUSULA 5 – DA CONFIDENCIALIDADE", 
     "5.1 As partes comprometem-se a manter sigilo sobre quaisquer informações confidenciais trocadas durante a execução do contrato."),
    ("CLÁUSULA 6 – DA RESCISÃO", 
     "6.1 O contrato poderá ser rescindido por qualquer das partes, mediante aviso prévio de 30 dias, por escrito.\n\n"
     "6.2 Em caso de rescisão, o CONTRATANTE pagará proporcionalmente pelos serviços já realizados até a data do aviso."),
    ("CLÁUSULA 7 – DO FORO", 
     "7.1 Para dirimir quaisquer controvérsias oriundas deste contrato, as partes elegem o foro da comarca de [cidade/estado], renunciando a qualquer outro, por mais privilegiado que seja."),
    ("", f"E, por estarem assim justos e contratados, firmam o presente instrumento, em duas vias de igual teor, na presença de testemunhas.\n\n[Local], {datetime.now().strftime('%d/%m/%Y')}"),
    ("CONTRATANTE:", "_"),
    ("CONTRATADO:", "_"),
    ("Testemunha 1:", "Nome:\nCPF:\nAssinatura:"),
    ("Testemunha 2:", "Nome:\nCPF:\nAssinatura:")
]

for titulo, conteudo in contrato_texto:
    if titulo:
        doc.add_heading(titulo, level=2)
    for linha in conteudo.split('\n'):
        doc.add_paragraph(linha)
git clone https://github.com/bealoreto/contrato.git
cd contrato
doc.save("Contrato_Criacao_Midias_Digitais.docx")
python -m venv venv
.\venv\Scripts\activate
